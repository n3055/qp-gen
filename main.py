import streamlit as st
import google.generativeai as genai
from docx import Document
from io import BytesIO

st.subheader("Question Paper Generator")
sub = st.selectbox("Select any subject", ("Operating System", "DSA", "Java"))
n = st.number_input("Number of question papers to generate", value=1)

if st.button("Generate QP", type="primary"):
    with st.spinner("Generating question papers..."):
        if sub == "Operating System":
            myfile = genai.upload_file("subjects/os.txt")
        elif sub == "DSA":
            myfile = genai.upload_file("subjects/dsa.txt")
        elif sub == "Java":
            myfile = genai.upload_file("subjects/JA.txt")

        # Configure the Generative AI API
        genai.configure(api_key=st.secrets["API_KEY"])
        model = genai.GenerativeModel("gemini-1.5-flash")
        papers = []

        for i in range(n):
            response = model.generate_content([
                myfile,
                """Generate a question paper based on the file provided.
                Question paper format:
                The Question paper for each course contains two parts, Part – A and Part – B. Part – A
                consists of 10 objective type questions for 20 marks (each question is worth 2 marks) covering the entire syllabus. Part – B
                Students have to answer five questions, one from each unit for 16 marks adding up to 80
                marks. Each main question may have a maximum of three sub-divisions. Each unit will have
                internal choice in which both questions cover the entire unit having the same complexity in terms of
                COs and Bloom’s taxonomy level."""
            ])
            papers.append(response.text)

    doc_files = []
    for i, paper in enumerate(papers):
        doc = Document()
        doc.add_heading(f"Question Paper {i + 1}", level=1)
        doc.add_paragraph(paper)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        doc_files.append((f"Question_Paper_{i + 1}.docx", buffer))

    for i, (filename, buffer) in enumerate(doc_files):
        with st.expander(f"Questions - {i + 1}"):
            st.download_button(
                label=f"Download {filename}",
                data=buffer,
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            st.write(papers[i])
           
